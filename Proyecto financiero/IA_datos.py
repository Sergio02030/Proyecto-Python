import matplotlib.pyplot as plt
from sklearn.linear_model import LinearRegression
import numpy as np
import math
from datetime import datetime
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

documento = Document()
titulo = documento.add_heading('RESULTADOS', 1)
titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

datos=int(input("Ingresa la cantidad de datos: "))
lista=[]
lista2=[]
for i in range(datos):
    dato=int(input("Ingresa el dato dependiente: "))
    dato2=int(input("Ingresa el dato independiente: "))
    lista.append(dato)
    lista2.append(dato2)

print(lista)
print(lista2)


def IA_validar(x,y):
    global pred
    X = np.array(x).reshape(-1, 1)
    Y = np.array(y)
    reg_1 = LinearRegression()
    reg_1.fit(X, Y)
    siguiente_x = np.array([[x[-1] + 1]])
    pred = reg_1.predict(siguiente_x)
    return print(f"Estimacion: {pred}")

IA_validar(lista2,lista)

def Regresion_lineal(x,y):
    global ecuacion,ecuacion2,ecuacion3,minimo,maximo,media_y,Desviacion_Estandar
    sumax=0; n=0; sumax2=0; sumay=0; sumaxy=0; sumay2=0; ecuacion=[]; ecuacion2=[]; ecuacion3=[]
    for i,j in zip(x,y):
        n+=1
        sumax2+= (i**2)
        sumax+=i
        sumay+=j
        sumay2+= (j**2)
        sumaxy+= (i*j)
    sumax_2= (sumax**2)
    b0 = ((sumay * sumax2) - (sumax * sumaxy)) / ((n * sumax2) - sumax_2)
    b1 = ((n * sumaxy) - (sumax * sumay)) / ((n * sumax2) - sumax_2)
    print(b0)
    print(b1)
    media_y = sumay / n
    Desviacion_Estandar = math.sqrt((n * sumay2 - (sumay ** 2)) / (n * (n - 1)))
    for t in x:
        e = b0 + (b1 * t)
        e1 = b0 + (b1 * t) - Desviacion_Estandar
        e2 = b0 + (b1 * t) + Desviacion_Estandar
        ecuacion.append(e)
        ecuacion2.append(e1)
        ecuacion3.append(e2)
    minimo=pred - Desviacion_Estandar
    maximo=pred + Desviacion_Estandar


def alerta(rango,desviacion):
    ahora = datetime.now()
    if nivel == "Alto":
        empresa = documento.add_paragraph(f"Anomalia Detectada - {ahora}")
        empresa.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        empresa = documento.add_paragraph(f"Valor esperado : {pred}")
        empresa.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        empresa = documento.add_paragraph(f"Valor observado : {rango}")
        empresa.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        empresa = documento.add_paragraph(f"Desviacion : +{desviacion}%")
        empresa.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    elif nivel == "Bajo":
        empresa = documento.add_paragraph(f"Anomalia Detectada - {ahora}")
        empresa.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        empresa = documento.add_paragraph(f"Valor esperado : {pred}")
        empresa.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        empresa = documento.add_paragraph(f"Valor observado : {rango}")
        empresa.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        empresa = documento.add_paragraph(f"Desviacion : -{desviacion}%")
        empresa.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    try:
        documento.save('REPORTE DEL DIA.docx')
    except PermissionError:
        print("Cerrar Word para actualizar")

    


Regresion_lineal(lista2,lista)

def ventana():
    fig,ax=plt.subplots()
    ax.scatter(lista2,lista)
    ax.scatter(lista2[-1]+1,pred)
    ax.plot(lista2,ecuacion)
    ax.plot(lista2,ecuacion2)
    ax.plot(lista2,ecuacion3)
    ax.set_title('Ventas vs Tiempo', loc = "center", fontdict = {'fontsize':14, 'fontweight':'bold', 'color':'tab:blue'})
    ax.set_ylabel("Ventas")
    ax.set_xlabel("Tiempo")
    plt.show()


ventana()

while True:
    dato_real=int(input("Ingresa el dato que salio: "))
    if dato_real > maximo:
        desviacion= ((dato_real - pred ) / pred) * 100
        nivel="Alto"
        alerta(dato_real,desviacion)
    elif dato_real < minimo:
        desviacion= ((dato_real - pred ) / pred) * 100
        nivel="Bajo"
        alerta(dato_real,desviacion)
    lista.append(dato_real)
    lista2.append(lista2[-1]+1)
    respuesta=input("Hay mas datos [Y/n]: ")
    if (respuesta == "Y"):
        continue
    elif (respuesta == "n"):
        break


print(lista2)
Regresion_lineal(lista2,lista)
ventana()


class Costo_de_Adquisición_Cliente:
    def __init__(self,inversion_total,clientes_generados):
        self.It=inversion_total
        self.Cg=clientes_generados
        
    def calculo_CAC(self):
        C=self.It / self.Cg
        print(f"El costo de adquisicion es de : $ {C}")

class Ticket_promedio:
    def __init__(self, Facturacion_total, Total_de_ventas):
        self.Factura = Facturacion_total
        self.ventas = Total_de_ventas

    def calculo(self):
        return self.Factura / self.ventas


ticket_anterior = int(input("Digite el ticket promedio del mes anterior: $"))

facturacion = int(input("Ingrese la facturación total: "))
ventas = int(input("Ingrese las ventas totales: "))

ticket_actual = Ticket_promedio(facturacion, ventas)

actual = ticket_actual.calculo()

print(f"Ticket promedio actual: ${actual:.2f}")

if actual > ticket_anterior:

    total = ((actual - ticket_anterior) / ticket_anterior) * 100

    print(f"Estado: Positivo")
    print(f"Incremento: +{total:.2f}%")

elif actual < ticket_anterior:

    total = ((ticket_anterior - actual) / ticket_anterior) * 100

    print(f"Estado: Atención")
    print(f"Disminución: -{total:.2f}%")

else:

    print("Estado: Estable")

class cumplimiento_de_meta:
    def __init__(self,ventas,meta):
        self.meta=meta
        self.ventas=ventas
        pass
    def cumplimiento(self):
        self.cumplimiento = (self.ventas / self.meta) * 100
        return self.cumplimiento

pregunta=input("Desea saber el cumplimiento de la meta [Y/n]: ")
if pregunta == "Y":
    v=int(input("Ingrese las ventas que obtuvo: $ "))
    m=int(input("Ingrese la meta de ventas que tenia propuesto: "))
    Cumplimiento_actual= cumplimiento_de_meta(v,m)
    cactual=Cumplimiento_actual.cumplimiento()
    print(f"El cumplimiento esta en un {cactual} %")
elif pregunta == "n":
    print("Gracias por usar el sistema inteligente, que tenga buena tarde")


    


    







